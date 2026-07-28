"""md_convert.docx_converter —— DOCX 转换器 (python-docx)。

将 markdown2 渲染后的 HTML 经 ``parse_html_to_blocks`` 解析为 Block 列表,
再映射为 python-docx 段落 / 表格 / 标题, 构建为 ``.docx`` 文档。

依赖: ``python-docx`` (可选, 仅在本模块被导入时加载)。

设计要点:
    - **InlineExtractor**: python-docx 不支持 HTML 内联格式, 需自行解析
      ``<strong>`` / ``<em>`` / ``<code>`` / ``<a>`` / ``<del>`` 为带
      bold / italic / code 属性的 Run。本模块用标准库 ``HTMLParser`` 实现。
    - 块级映射: heading→Heading 样式段落, paragraph→Normal 段落,
      list→List Bullet / List Number 样式, code→等宽字体段落+灰底,
      quote→缩进段落, table→Table Grid 样式表格, hr→底边框段落。
"""

from __future__ import annotations

from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from agentkit.assets import CODE_FONT, DEFAULT_FONT
from agentkit.tools.md_convert.base import Block, Converter, parse_html_to_blocks


# ---------------------------------------------------------------------------
# InlineExtractor —— 内联 HTML → Run 列表
# ---------------------------------------------------------------------------
@dataclass
class _RunSpec:
    """内联格式的单个 run 描述。

    Attributes:
        text:   run 文本 (已解码, 无 HTML 标签)。
        bold:   是否粗体。
        italic: 是否斜体。
        code:   是否行内代码 (等宽字体)。
        strike: 是否删除线。
    """

    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False
    strike: bool = False


class InlineExtractor(HTMLParser):
    """解析内联 HTML, 输出 ``_RunSpec`` 列表。

    追踪 ``<strong>`` / ``<em>`` / ``<code>`` / ``<del>`` 的嵌套状态,
    每个 ``handle_data`` 产生一个 ``_RunSpec``。连续相同格式的 run 不合并
    (正确性优先, 合并由 python-docx 内部处理或后续优化)。

    超链接 (``<a>``) 仅提取文本, 不创建实际超链接 (python-docx 需 XML 操作,
    超出当前范围)。
    """

    _BOLD_TAGS = frozenset({"strong", "b"})
    _ITALIC_TAGS = frozenset({"em", "i"})
    _STRIKE_TAGS = frozenset({"del", "strike", "s"})

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.runs: list[_RunSpec] = []
        self._bold = False
        self._italic = False
        self._code = False
        self._strike = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self._BOLD_TAGS:
            self._bold = True
        elif tag in self._ITALIC_TAGS:
            self._italic = True
        elif tag == "code":
            self._code = True
        elif tag in self._STRIKE_TAGS:
            self._strike = True
        # <a>, <span>, <font> 等: 仅跟踪文本, 不改变格式

    def handle_endtag(self, tag: str) -> None:
        if tag in self._BOLD_TAGS:
            self._bold = False
        elif tag in self._ITALIC_TAGS:
            self._italic = False
        elif tag == "code":
            self._code = False
        elif tag in self._STRIKE_TAGS:
            self._strike = False

    def handle_startendtag(self, tag: str,
                           attrs: list[tuple[str, str | None]]) -> None:
        if tag == "br":
            self.runs.append(_RunSpec(text="\n"))

    def handle_data(self, data: str) -> None:
        if not data:
            return
        self.runs.append(_RunSpec(
            text=data,
            bold=self._bold,
            italic=self._italic,
            code=self._code,
            strike=self._strike,
        ))


def _extract_runs(html: str) -> list[_RunSpec]:
    """便捷函数: 解析内联 HTML, 返回 ``_RunSpec`` 列表。"""
    extractor = InlineExtractor()
    extractor.feed(html)
    extractor.close()
    return extractor.runs


def _add_runs_to_paragraph(paragraph, html: str) -> None:
    """将内联 HTML 解析为 runs 并添加到 python-docx 段落。

    Args:
        paragraph: ``docx.text.paragraph.Paragraph`` 实例。
        html:      内联 HTML 字符串 (如 ``"text <strong>bold</strong>"``)。
    """
    for spec in _extract_runs(html):
        run = paragraph.add_run(spec.text)
        if spec.bold:
            run.bold = True
        if spec.italic:
            run.italic = True
        if spec.strike:
            run.font.strike = True
        if spec.code:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)
            # 确保东亚字符也使用等宽字体
            run._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)


# ---------------------------------------------------------------------------
# OOXML 辅助: 段落底色 / 水平线
# ---------------------------------------------------------------------------
def _set_paragraph_shading(paragraph, fill: str) -> None:
    """设置段落背景色 (OOXML ``w:shd``)。

    Args:
        paragraph: python-docx Paragraph 实例。
        fill:      6 位 hex 颜色 (不含 ``#``), 如 ``"f6f8fa"``。
    """
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _add_horizontal_rule(doc) -> None:
    """向文档添加水平线段落 (底边框)。"""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ---------------------------------------------------------------------------
# 文档默认字体配置
# ---------------------------------------------------------------------------
def _apply_font_to_style(style, font_name: str) -> None:
    """将 ``font_name`` 同时设为样式的 Latin (ascii/hAnsi) 与 CJK (eastAsia) 字体。

    python-docx 的 ``style.font.name`` 仅设置 ``w:ascii`` / ``w:hAnsi``,
    东亚字符仍会回退到系统默认字体。本函数补设 ``w:eastAsia`` 与 ``w:cs``
    (复杂文种), 确保 OPPO Sans 覆盖中英文。

    Args:
        style:      ``docx.styles.style.BaseStyle`` 实例 (如 Normal / Heading 1)。
        font_name:  字体名 (如 ``"OPPO Sans"``)。
    """
    style.font.name = font_name
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _configure_document_fonts(doc) -> None:
    """将文档默认字体 (Normal) 及所有 Heading 样式设为 OPPO Sans。

    List Bullet / List Number 等样式基于 Normal 继承, 无需单独设置。
    Heading 样式在默认模板中有显式字体, 必须逐个覆盖。
    """
    normal = doc.styles["Normal"]
    _apply_font_to_style(normal, DEFAULT_FONT)
    for level in range(1, 10):
        name = f"Heading {level}"
        if name in [s.name for s in doc.styles]:
            _apply_font_to_style(doc.styles[name], DEFAULT_FONT)


# ---------------------------------------------------------------------------
# DocxConverter
# ---------------------------------------------------------------------------
class DocxConverter(Converter):
    """将 HTML 转换为 DOCX 文档 (基于 python-docx)。

    流程: HTML → Block 列表 → python-docx 段落/表格 → Document.save。
    """

    @property
    def format(self) -> str:
        return "docx"

    def convert(self, html: str, output_path: Path) -> None:
        """构建 DOCX 并写入 ``output_path``。

        Args:
            html:        markdown2 渲染后的 HTML 片段。
            output_path: ``.docx`` 文件路径。
        """
        blocks = parse_html_to_blocks(html)
        doc = Document()
        _configure_document_fonts(doc)
        for block in blocks:
            _render_block(doc, block)
        doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Block → python-docx 渲染
# ---------------------------------------------------------------------------
def _render_block(doc, block: Block) -> None:
    """将单个 Block 渲染为 python-docx 元素并追加到 ``doc``。"""
    match block.type:
        case "heading":
            _render_heading(doc, block)
        case "paragraph":
            p = doc.add_paragraph()
            _add_runs_to_paragraph(p, block.text)
        case "list":
            _render_list(doc, block)
        case "code":
            _render_code(doc, block)
        case "quote":
            p = doc.add_paragraph()
            _add_runs_to_paragraph(p, block.text)
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.right_indent = Pt(10)
            _set_paragraph_shading(p, "f6f8fa")
        case "table":
            _render_table(doc, block)
        case "hr":
            _add_horizontal_rule(doc)


def _render_heading(doc, block: Block) -> None:
    """渲染标题段落 (带内联格式)。"""
    level = min(max(block.level, 1), 9)
    p = doc.add_paragraph(style=f"Heading {level}")
    _add_runs_to_paragraph(p, block.text)


def _render_list(doc, block: Block) -> None:
    """渲染列表项 (无序用 List Bullet, 有序用 List Number)。"""
    style = "List Number" if block.ordered else "List Bullet"
    for item in block.items:
        p = doc.add_paragraph(style=style)
        _add_runs_to_paragraph(p, item)


def _render_code(doc, block: Block) -> None:
    """渲染代码块 (等宽字体 + 灰色背景)。"""
    p = doc.add_paragraph()
    run = p.add_run(block.text)
    run.font.name = CODE_FONT
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)
    p.paragraph_format.left_indent = Pt(10)
    _set_paragraph_shading(p, "f6f8fa")


def _render_table(doc, block: Block) -> None:
    """渲染表格 (Table Grid 样式, 表头加粗)。"""
    headers = block.headers
    rows = block.rows
    n_cols = max(len(headers), max((len(r) for r in rows), default=0)) or 1

    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # 表头行
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = header_cells[i].paragraphs[0]
        _add_runs_to_paragraph(p, header)
        for run in p.runs:
            run.bold = True
    # 补齐空表头
    for i in range(len(headers), n_cols):
        header_cells[i].paragraphs[0].add_run("")

    # 数据行
    for r_idx, row in enumerate(rows, 1):
        cells = table.rows[r_idx].cells
        for c_idx, cell_text in enumerate(row):
            p = cells[c_idx].paragraphs[0]
            _add_runs_to_paragraph(p, cell_text)
        for c_idx in range(len(row), n_cols):
            cells[c_idx].paragraphs[0].add_run("")
